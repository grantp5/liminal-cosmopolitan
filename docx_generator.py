import docx
from docx.shared import Pt

def create_docx(final_order_tu, final_order_b, name):
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    for x in final_order_tu:
        number = final_order_tu.index(x ) +1
        a = doc.add_paragraph(str(number))
        a.add_run('. ')
        b = doc.add_paragraph('ANSWER: ')
        c = b.add_run('ANSWER')
        c.bold = True
        c.underline = True
        d = doc.add_paragraph('<')
        d.add_run(x)
        d.add_run('>')
        doc.add_paragraph()

    doc.add_page_break()

    for x in final_order_b:
        number = final_order_b.index(x ) +1
        a = doc.add_paragraph(str(number))
        a.add_run('. For 10 points each:')
        for y in range(3):
            b = doc.add_paragraph('[10] ')
            c = doc.add_paragraph('ANSWER: ')
            d = c.add_run('ANSWER')
            d.bold = True
            d.underline = True
        d = doc.add_paragraph('<')
        d.add_run(x)
        d.add_run('>')
        doc.add_paragraph()

    doc.save(name)

    print('Packet Generated!')